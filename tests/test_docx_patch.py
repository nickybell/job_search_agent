"""Applying the structured tailoring patch to a .docx (Step 4's mechanism).

NOTE: agent-authored and awaiting Nicky's review (the standing rule in
TODO.md) -- one behaviour per test, names state the claim.

The load-bearing claims: the numbered text the model sees maps exactly onto
the paragraphs the patch targets; applying a patch changes only the targeted
text and keeps run formatting; and a patch targeting anything never offered
fails loudly instead of guessing.
"""

from __future__ import annotations

import pytest
from docx import Document

from jsa.docx_patch import (
    PatchChange,
    PatchError,
    TailoringPatch,
    apply_patch,
    iter_paragraphs,
    parse_patch,
    render_changelog,
    render_numbered_text,
)


def make_doc() -> Document:
    # A name line, an empty spacer, a bold bullet, a plain bullet.
    doc = Document()
    doc.add_paragraph("Nicholas Bell")
    doc.add_paragraph("")
    bold = doc.add_paragraph()
    bold.add_run("Led onboarding programs").bold = True
    doc.add_paragraph("Wrote help-center articles")
    return doc


def index_of(doc, text: str) -> int:
    # Indexes are derived, not hardcoded, so the tests don't depend on how
    # many paragraphs python-docx's default template starts with.
    return next(i for i, p in enumerate(iter_paragraphs(doc)) if p.text == text)


def patch_of(*changes: dict) -> TailoringPatch:
    return TailoringPatch(base="template", changes=[PatchChange(**c) for c in changes])


# --- the numbered view the model sees --------------------------------------


def test_numbered_ids_map_onto_paragraph_order_and_spacers_are_not_offered():
    doc = make_doc()
    text = render_numbered_text(doc)
    for i, paragraph in enumerate(iter_paragraphs(doc)):
        if paragraph.text.strip():
            assert f"[P{i}] {paragraph.text.strip()}" in text
        else:
            assert f"[P{i}]" not in text


def test_table_cell_paragraphs_are_numbered_and_patchable():
    # Resumes are frequently laid out in tables, which document.paragraphs
    # does not descend into -- the walker must.
    doc = Document()
    doc.add_paragraph("Body paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].text = "Left cell"
    table.cell(0, 1).paragraphs[0].text = "Right cell"

    idx = index_of(doc, "Left cell")
    assert f"[P{idx}] Left cell" in render_numbered_text(doc)

    apply_patch(doc, patch_of({"paragraph": idx, "replacement": "New left", "rationale": "r"}))
    assert table.cell(0, 0).paragraphs[0].text == "New left"


# --- applying the patch -----------------------------------------------------


def test_apply_replaces_only_the_targeted_paragraph_and_reports_it():
    doc = make_doc()
    idx = index_of(doc, "Wrote help-center articles")
    applied = apply_patch(
        doc,
        patch_of(
            {"paragraph": idx, "replacement": "Wrote adoption playbooks", "rationale": "JD fit"}
        ),
    )
    texts = [p.text for p in iter_paragraphs(doc)]
    assert "Wrote adoption playbooks" in texts
    assert "Wrote help-center articles" not in texts
    assert "Nicholas Bell" in texts  # untargeted paragraphs untouched
    assert len(applied) == 1
    assert (applied[0].before, applied[0].after) == (
        "Wrote help-center articles",
        "Wrote adoption playbooks",
    )
    assert applied[0].rationale == "JD fit"


def test_apply_keeps_the_leading_runs_character_formatting():
    # Assigning paragraph.text would reset a bold heading to the style
    # default; the replacement must keep the leading run's formatting.
    doc = make_doc()
    idx = index_of(doc, "Led onboarding programs")
    apply_patch(
        doc, patch_of({"paragraph": idx, "replacement": "Led enablement", "rationale": "r"})
    )
    target = list(iter_paragraphs(doc))[idx]
    assert target.text == "Led enablement"
    assert target.runs[0].bold is True


def test_an_out_of_range_target_raises():
    with pytest.raises(PatchError):
        apply_patch(make_doc(), patch_of({"paragraph": 999, "replacement": "x", "rationale": "r"}))


def test_an_empty_spacer_target_raises():
    # Spacers are never offered in the numbered text, so a patch that targets
    # one is malformed, not a judgment call.
    doc = make_doc()
    idx = index_of(doc, "")
    with pytest.raises(PatchError):
        apply_patch(doc, patch_of({"paragraph": idx, "replacement": "x", "rationale": "r"}))


def test_a_no_op_change_is_dropped_not_applied():
    doc = make_doc()
    idx = index_of(doc, "Nicholas Bell")
    applied = apply_patch(
        doc, patch_of({"paragraph": idx, "replacement": "Nicholas Bell", "rationale": "r"})
    )
    assert applied == []


# --- parsing the model's output ---------------------------------------------


def test_parse_patch_tolerates_fences_and_surrounding_prose():
    raw = (
        "Here is the patch:\n```json\n"
        '{"base": "customer-education", "summary": "s", "changes": []}\n'
        "```\nDone."
    )
    patch = parse_patch(raw)
    assert patch.base == "customer-education"
    assert patch.summary == "s"
    assert patch.new_family is None
    assert patch.changes == []


def test_parse_patch_rejects_non_json_output():
    with pytest.raises(PatchError):
        parse_patch("I would rather answer in prose.")


def test_parse_patch_rejects_a_patch_without_a_base():
    # base is how the patch names its template; a patch that omits it cannot
    # be applied to anything and must fail loudly, not guess.
    with pytest.raises(PatchError):
        parse_patch('{"summary": "s", "changes": []}')


# --- the changelog ----------------------------------------------------------


def test_changelog_renders_each_change_with_before_after_and_rationale():
    doc = make_doc()
    idx = index_of(doc, "Wrote help-center articles")
    applied = apply_patch(
        doc, patch_of({"paragraph": idx, "replacement": "Wrote playbooks", "rationale": "JD fit"})
    )
    text = render_changelog(
        company="Acme",
        title="Enablement Lead",
        model="claude-opus-4-8",
        date_generated="2026-08-21",
        base="customer-education",
        base_rationale="closest role family",
        new_template=None,
        summary="Emphasize enablement.",
        applied=applied,
    )
    assert "Acme" in text and "Enablement Lead" in text
    assert "`customer-education` template" in text
    assert "closest role family" in text
    assert "**Before:** Wrote help-center articles" in text
    assert "**After:** Wrote playbooks" in text
    assert "**Why:** JD fit" in text


def test_changelog_flags_a_new_template_for_curation():
    text = render_changelog(
        company="Acme",
        title="AI Enablement Lead",
        model="claude-opus-4-8",
        date_generated="2026-08-22",
        base="customer-education",
        base_rationale=None,
        new_template="ai-enablement",
        summary=None,
        applied=[],
    )
    assert "NEW TEMPLATE CREATED" in text
    assert "resume_templates/ai-enablement.docx" in text


def test_changelog_says_so_when_nothing_changed():
    text = render_changelog(
        company="Acme",
        title="Enablement Lead",
        model="claude-opus-4-8",
        date_generated="2026-08-21",
        base="customer-education",
        base_rationale=None,
        new_template=None,
        summary=None,
        applied=[],
    )
    assert "no changes" in text.lower()
