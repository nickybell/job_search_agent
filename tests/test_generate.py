"""jsa generate: the Step 4 tailoring pass and its seam into Step 5.

NOTE: agent-authored and awaiting Nicky's review (the standing rule in
TODO.md).

Hermetic: a throwaway ``file:`` SQLite DB, a stub tailor callable instead of
the model, a stub soffice that just creates the .pdf, a stub gws binary for
the tracker append, and a tmp-path template library -- nothing touches the
network, LibreOffice, a real Sheet, or the real resume_templates/.
"""

from __future__ import annotations

import json

import pytest
from conftest import make_posting
from docx import Document

from jsa import db, generate

EDU_BULLET = "Designed role-based learning paths for enterprise customers"
CS_BULLET = "Drove adoption, retention, and expansion for enterprise accounts"


@pytest.fixture(autouse=True)
def packets_root(tmp_path, monkeypatch):
    root = tmp_path / "Job Applications"
    monkeypatch.setenv("JSA_PACKETS_DIR", str(root))
    return root


@pytest.fixture(autouse=True)
def templates_dir(tmp_path, monkeypatch):
    """A two-template library standing in for resume_templates/."""
    directory = tmp_path / "resume_templates"
    directory.mkdir()
    for slug, bullet in (("customer-education", EDU_BULLET), ("customer-success", CS_BULLET)):
        doc = Document()
        doc.add_paragraph("Nicholas Bell")
        doc.add_paragraph(bullet)
        doc.save(str(directory / f"{slug}.docx"))
    monkeypatch.setenv("JSA_RESUME_TEMPLATES_DIR", str(directory))
    return directory


@pytest.fixture(autouse=True)
def soffice_stub(tmp_path, monkeypatch):
    """Stands in for LibreOffice: creates the sibling .pdf and exits 0."""
    script = tmp_path / "soffice-stub"
    script.write_text(
        "#!/bin/sh\n"
        "# argv: --headless --convert-to pdf --outdir <dir> <docx>\n"
        'out="$5"\n'
        'in="$6"\n'
        'base=$(basename "$in" .docx)\n'
        'touch "$out/$base.pdf"\n'
    )
    script.chmod(0o755)
    monkeypatch.setenv("JSA_SOFFICE_BIN", str(script))
    return script


@pytest.fixture(autouse=True)
def gws_calls(tmp_path, monkeypatch):
    """A gws stub that confirms every append; returns the invocation log path."""
    calls = tmp_path / "gws-calls.log"
    script = tmp_path / "gws-stub"
    script.write_text(
        f'#!/bin/sh\necho "$@" >> "{calls}"\necho \'{{"updates":{{"updatedRows":1}}}}\'\n'
    )
    script.chmod(0o755)
    monkeypatch.setenv("JSA_GWS_BIN", str(script))
    return calls


def seed(client, *, decision="Apply", tracked=0, jd="# The JD", n=0) -> int:
    posting_id = db.insert_posting(
        client,
        make_posting(
            company=f"Acme{n or ''}",
            normalized_company=f"Acme{n or ''}",
            url=f"https://job-boards.greenhouse.io/acme/jobs/{100 + n}",
            canonical_url=f"https://job-boards.greenhouse.io/acme/jobs/{100 + n}",
        ),
    )
    if jd is not None:
        db.update_jd_capture(client, posting_id, jd_markdown=jd, location="Remote", title=None)
    if decision:
        db.record_decision(client, posting_id, decision, None)
    if tracked:
        db.mark_tracked(client, posting_id)
    return posting_id


def make_tailor(
    base="customer-education",
    changes=None,
    *,
    new_family=None,
    summary="Emphasize enablement.",
):
    """A tailor callable returning a fixed patch; records the prompts it saw."""
    payload = {
        "base": base,
        "base_rationale": "closest role family",
        "new_family": new_family,
        "summary": summary,
        "changes": changes or [],
    }

    def fake(prompt: str) -> str:
        fake.calls.append(prompt)
        return json.dumps(payload)

    fake.calls = []
    return fake


def bullet_index(templates_dir, slug: str, text: str) -> int:
    # Derived, not hardcoded, so the tests don't depend on the default
    # template's initial paragraph count.
    doc = Document(str(templates_dir / f"{slug}.docx"))
    return next(i for i, p in enumerate(doc.paragraphs) if p.text == text)


def tracked_flag(client, posting_id) -> int:
    cursor = client.execute("SELECT added_to_tracker FROM postings WHERE id = ?", (posting_id,))
    return int(cursor.fetchone()[0])


# --- the happy path ---------------------------------------------------------


def test_tailors_renders_both_formats_and_tracks(config, client, packets_root, templates_dir):
    posting_id = seed(client)
    replacement = "Designed enterprise onboarding and enablement programs"
    tailor = make_tailor(
        changes=[
            {
                "paragraph": bullet_index(templates_dir, "customer-education", EDU_BULLET),
                "replacement": replacement,
                "rationale": "match the JD",
            }
        ]
    )

    summary, results = generate.run_generate(config, tailor=tailor)

    assert (summary.eligible, summary.generated, summary.tracked) == (1, 1, 1)
    path = packets_root / "Acme - Customer Enablement Lead"
    assert results[0].path == path
    assert results[0].base == "customer-education"
    assert (path / "job_posting.md").read_text() == "# The JD\n"
    # prd: spaces removed from the file names, kept in the directory name.
    docx_file = path / "NicholasBell_Resume_CustomerEnablementLead_Acme.docx"
    assert docx_file.is_file()
    assert (path / "NicholasBell_Resume_CustomerEnablementLead_Acme.pdf").is_file()
    assert replacement in [p.text for p in Document(str(docx_file)).paragraphs]
    changelog = (path / "resume_changelog.md").read_text()
    assert "match the JD" in changelog
    # The template pick and its rationale are part of what the user reviews.
    assert "`customer-education` template" in changelog
    assert "closest role family" in changelog
    # The seam into Step 5: flag set only after the (stubbed) API confirmation.
    assert tracked_flag(client, posting_id) == 1


def test_the_prompt_carries_the_jd_and_every_template(config, client):
    seed(client)
    tailor = make_tailor()
    generate.run_generate(config, tailor=tailor)
    prompt = tailor.calls[0]
    assert "# The JD" in prompt
    assert "### Template: customer-education" in prompt
    assert "### Template: customer-success" in prompt
    assert EDU_BULLET in prompt and CS_BULLET in prompt
    assert "{{" not in prompt  # every template slot interpolated


# --- the template library ---------------------------------------------------


def test_an_unknown_base_template_fails_the_row(config, client):
    posting_id = seed(client)

    summary, results = generate.run_generate(config, tailor=make_tailor(base="project-management"))

    assert (summary.failed, summary.generated) == (1, 0)
    assert "unknown template" in results[0].error
    assert tracked_flag(client, posting_id) == 0


def test_new_family_seeds_a_template_and_flags_the_changelog(
    config, client, packets_root, templates_dir
):
    # The outward-expansion rule: no family fits, so the nearest template is
    # the base and the tailored result becomes the new family's template,
    # flagged for the curation scrub.
    seed(client)

    summary, results = generate.run_generate(config, tailor=make_tailor(new_family="AI Enablement"))

    assert summary.generated == 1
    assert results[0].new_template == "ai-enablement"
    assert (templates_dir / "ai-enablement.docx").is_file()
    changelog = (results[0].path / "resume_changelog.md").read_text()
    assert "NEW TEMPLATE CREATED" in changelog
    assert "resume_templates/ai-enablement.docx" in changelog


def test_new_family_never_clobbers_an_existing_template(config, client, templates_dir):
    # A colliding declaration means the family already exists; curated work is
    # never overwritten, and the row still generates normally.
    seed(client)
    before = (templates_dir / "customer-success.docx").read_bytes()

    summary, results = generate.run_generate(
        config, tailor=make_tailor(new_family="customer-success")
    )

    assert summary.generated == 1
    assert results[0].new_template is None
    assert (templates_dir / "customer-success.docx").read_bytes() == before


def test_an_empty_template_library_fails_loudly_before_any_row(
    config, client, tmp_path, monkeypatch
):
    seed(client)
    empty = tmp_path / "empty-library"
    empty.mkdir()
    monkeypatch.setenv("JSA_RESUME_TEMPLATES_DIR", str(empty))
    with pytest.raises(generate.GenerateError):
        generate.run_generate(config, tailor=make_tailor())


# --- ensure-semantics (contrast with jsa packet's standalone skip) ----------


def test_a_bare_directory_is_reentered_and_completed(config, client, packets_root):
    # The completion guard is added_to_tracker, not directory-exists: a bare
    # directory (interrupted run, refetch rebuild) is resumable work.
    seed(client)
    bare = packets_root / "Acme - Customer Enablement Lead"
    bare.mkdir(parents=True)

    summary, _ = generate.run_generate(config, tailor=make_tailor())

    assert summary.generated == 1
    assert (bare / "job_posting.md").is_file()
    assert (bare / "resume_changelog.md").is_file()


# --- the NULL-JD rule -------------------------------------------------------


def test_a_null_jd_row_is_skipped_never_tailored_blind(config, client, gws_calls):
    posting_id = seed(client, jd=None)
    tailor = make_tailor()

    summary, results = generate.run_generate(config, tailor=tailor)

    assert (summary.generated, summary.skipped_no_jd) == (0, 1)
    assert results[0].status == "skipped_no_jd"
    assert tailor.calls == []  # the model is never asked
    assert results[0].path.is_dir()  # the directory is still ensured
    assert not gws_calls.exists()  # and the tracker call is skipped
    assert tracked_flag(client, posting_id) == 0  # the row stays in the queue


def test_a_hand_filled_job_posting_is_used_and_never_overwritten(config, client, packets_root):
    # The escape hatch for rows refetch can never reach (unsupported ATS):
    # a hand-pasted job_posting.md is input, not output.
    seed(client, jd=None)
    path = packets_root / "Acme - Customer Enablement Lead"
    path.mkdir(parents=True)
    hand_filled = "# Hand-pasted JD\nA Workday role.\n"
    (path / "job_posting.md").write_text(hand_filled)
    tailor = make_tailor()

    summary, results = generate.run_generate(config, tailor=tailor)

    assert summary.generated == 1
    assert results[0].jd_source == "packet"
    assert "Hand-pasted JD" in tailor.calls[0]
    assert (path / "job_posting.md").read_text() == hand_filled


# --- eligibility and the Step 5 seam ----------------------------------------


def test_default_queue_excludes_tracked_rows(config, client):
    seed(client, tracked=1)
    summary, _ = generate.run_generate(config, tailor=make_tailor())
    assert summary.eligible == 0


def test_id_waives_the_tracker_condition_and_the_append_no_ops(config, client, gws_calls):
    # The refetch-regenerate path: a tracked-but-unapplied row is rebuilt via
    # --id, and the closing track call must not append it a second time.
    posting_id = seed(client, tracked=1)

    summary, results = generate.run_generate(config, posting_id=posting_id, tailor=make_tailor())

    assert summary.generated == 1
    assert results[0].already_tracked is True
    assert results[0].tracked is False
    assert not gws_calls.exists()  # no second Sheet append


def test_a_failed_append_keeps_the_flag_unset_and_reports_it(config, client, tmp_path, monkeypatch):
    posting_id = seed(client)
    bad = tmp_path / "gws-bad"
    bad.write_text("#!/bin/sh\nexit 1\n")
    bad.chmod(0o755)
    monkeypatch.setenv("JSA_GWS_BIN", str(bad))

    summary, results = generate.run_generate(config, tailor=make_tailor())

    assert summary.generated == 1  # the resume work still stands
    assert summary.track_failed == 1
    assert results[0].track_error
    assert tracked_flag(client, posting_id) == 0  # stays in the Step 5 backlog


# --- failure and preview ----------------------------------------------------


def test_a_malformed_patch_fails_the_row_and_leaves_it_queued(config, client):
    posting_id = seed(client)

    summary, results = generate.run_generate(
        config, tailor=lambda prompt: "I would rather answer in prose."
    )

    assert (summary.failed, summary.generated) == (1, 0)
    assert results[0].status == "failed"
    assert tracked_flag(client, posting_id) == 0
    assert [r[0] for r in db.pending_packets(client)] == [posting_id]  # still queued


def test_dry_run_touches_nothing(config, client, packets_root, gws_calls):
    seed(client)
    tailor = make_tailor()

    summary, results = generate.run_generate(config, dry_run=True, tailor=tailor)

    assert results[0].status == "would_generate"
    assert tailor.calls == []
    assert not packets_root.exists()
    assert not gws_calls.exists()
